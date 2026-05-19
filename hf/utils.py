"""
Pure utility functions and backend HTTP client for the RAG frontend.
This module deliberately has no Gradio dependency so it can be unit-tested
in isolation (and run on Python versions where Gradio's dep chain is broken).
"""
import os
import json
import requests
from typing import Optional, Generator, Tuple


# ---------- Configuration Constants ----------

BACKEND_URLS = {
    "query":  "https://carsonbytes--query.modal.run/",
    "debug":  "https://carsonbytes--debug.modal.run/",
    "upload": "https://carsonbytes--upload.modal.run/",
    "health": "https://carsonbytes--health.modal.run/",
    "list":   "https://carsonbytes--list.modal.run/",
    "delete": "https://carsonbytes--delete.modal.run/",
}

WARMUP_TIMEOUT = 15
QUERY_TIMEOUT = 60
UPLOAD_TIMEOUT = 120
MAX_FILE_SIZE = 10 * 1024 * 1024
MAX_QUESTION_LENGTH = 1000


# ---------- Sample Data ----------

def load_sample_data() -> dict:
    """Load TechVision Corp sample annual reports from JSON."""
    data_path = os.path.join(os.path.dirname(__file__), "data", "sample_reports.json")
    with open(data_path, "r", encoding="utf-8") as f:
        return json.load(f)


# ---------- Pure utility functions ----------

def get_backend_url(endpoint: str) -> str:
    return BACKEND_URLS.get(endpoint, BACKEND_URLS["query"])


def format_file_size(size_bytes: int) -> str:
    if size_bytes < 1024:
        return f"{size_bytes} bytes"
    elif size_bytes < 1024 * 1024:
        return f"{size_bytes / 1024:.1f} KB"
    else:
        return f"{size_bytes / (1024 * 1024):.1f} MB"


def clean_answer(answer: str) -> str:
    """Strip leaked prompt-template markers from LLM output."""
    for marker in ['[/INST]', '[INST]', '<<SYS>>', '<</SYS>>', '<</SYS', '</s>', '<|im_end|>', '<|im_start|>']:
        if marker in answer:
            answer = answer[:answer.index(marker)]
    return answer.strip()


def update_index_display(docs: list) -> str:
    """Human-readable status for the index display markdown component."""
    if not docs:
        return "📄 *No documents indexed yet*"
    count = len(docs)
    names = ", ".join(f"`{d}`" for d in docs[:3])
    suffix = f" (and {count - 3} more)" if count > 3 else ""
    return f"📚 **{count} document{'s' if count != 1 else ''} indexed:** {names}{suffix}"


# ---------- File extraction ----------

def _extract_pdf_text(path: str) -> str:
    from pypdf import PdfReader
    reader = PdfReader(path)
    parts = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception:
            continue
    return "\n\n".join(parts)


def _extract_docx_text(path: str) -> str:
    import docx
    doc = docx.Document(path)
    paragraphs = [p.text for p in doc.paragraphs if p.text]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    paragraphs.append(cell.text)
    return "\n\n".join(paragraphs)


def extract_text_from_file(path: str) -> Tuple[Optional[str], Optional[str]]:
    """Return (text, error). Supports .txt, .pdf, .docx."""
    ext = os.path.splitext(path)[1].lower()
    try:
        if ext == ".txt":
            with open(path, "r", encoding="utf-8") as f:
                return f.read(), None
        if ext == ".pdf":
            return _extract_pdf_text(path), None
        if ext == ".docx":
            return _extract_docx_text(path), None
        return None, f"❌ Unsupported file type: {ext}. Use .txt, .pdf, or .docx."
    except UnicodeDecodeError:
        return None, "❌ Could not decode .txt file. Please use UTF-8 encoding."
    except Exception as e:
        return None, f"❌ Failed to read {ext} file: {e}"


def resolve_file_path(file_obj) -> Optional[str]:
    """gr.File may return a str path (Gradio 6+) or an object with .name (older)."""
    if file_obj is None:
        return None
    if isinstance(file_obj, str):
        return file_obj
    return getattr(file_obj, "name", None) or getattr(file_obj, "path", None)


# ---------- Backend HTTP client ----------

def check_health() -> Tuple[bool, list]:
    """GET /health. Returns (index_exists, docs_list)."""
    try:
        resp = requests.get(get_backend_url("health"), timeout=WARMUP_TIMEOUT)
        if resp.status_code == 200:
            data = resp.json()
            docs = data.get("indexed_docs", [])
            return data.get("index_exists", len(docs) > 0), docs
        return False, []
    except Exception:
        return False, []


def fetch_indexed_docs() -> list:
    """GET /list. Returns list of indexed doc IDs (empty on error)."""
    try:
        resp = requests.get(get_backend_url("list"), timeout=WARMUP_TIMEOUT)
        if resp.status_code == 200:
            return resp.json().get("docs", [])
        return []
    except Exception:
        return []


def upload_text_to_backend(text_content: str, filename: str) -> Tuple[bool, Optional[str], Optional[int], str]:
    """
    POST /upload. Returns (success, filename, char_count, message).
    Returns the ORIGINAL filename (matches /list output) so the dropdown can
    select the freshly uploaded doc without a name mismatch.
    """
    payload = {"text": text_content, "filename": filename}
    try:
        response = requests.post(
            get_backend_url("upload"),
            json=payload,
            timeout=UPLOAD_TIMEOUT,
        )
        response.raise_for_status()
        data = response.json()
        if data.get("status") == "success":
            returned_name = data.get("filename") or filename
            char_count = data.get("char_count")
            return True, returned_name, char_count, f"✅ Indexed **{returned_name}**"
        return False, None, None, f"❌ {data.get('message', 'Upload failed')}"
    except requests.exceptions.Timeout:
        return False, None, None, "❌ Upload timed out. Please try a smaller file."
    except Exception as e:
        return False, None, None, f"❌ Upload failed: {str(e)}"


def delete_document(filename: str) -> Tuple[bool, list, str]:
    """
    POST /delete. Removes the named document from the index.
    Returns (success, remaining_docs, message).
    """
    if not filename:
        return False, [], "⚠️ No document selected."
    try:
        response = requests.post(
            get_backend_url("delete"),
            json={"filename": filename},
            timeout=UPLOAD_TIMEOUT,
        )
        response.raise_for_status()
        data = response.json()
        if data.get("status") == "success":
            return True, data.get("indexed_docs", []), f"🗑️ Removed **{filename}**"
        if data.get("status") == "not_found":
            return False, data.get("indexed_docs", []), f"⚠️ '{filename}' not found."
        return False, [], f"❌ {data.get('message', 'Delete failed')}"
    except requests.exceptions.Timeout:
        return False, [], "❌ Delete timed out."
    except Exception as e:
        return False, [], f"❌ Delete failed: {str(e)}"


def upload_pasted_text(text: str, filename: str = "pasted.txt") -> Tuple[bool, Optional[str], Optional[int], str]:
    """Validate and upload pasted text directly."""
    if not text or not text.strip():
        return False, None, None, "⚠️ Please paste some text first."
    filename = (filename or "pasted.txt").strip() or "pasted.txt"
    if not filename.endswith(".txt"):
        filename += ".txt"
    return upload_text_to_backend(text, filename)


def upload_document(file_obj, max_file_size: int = MAX_FILE_SIZE) -> Tuple[bool, Optional[str], Optional[int], str]:
    """Parse a file object (path or Gradio file) and upload its extracted text."""
    path = resolve_file_path(file_obj)
    if not path:
        return False, None, None, "❌ No file selected."
    if not os.path.exists(path):
        return False, None, None, f"❌ File not found: {path}"
    file_size = os.path.getsize(path)
    if file_size > max_file_size:
        return False, None, None, f"❌ File too large. Max {format_file_size(max_file_size)}."
    text_content, err = extract_text_from_file(path)
    if err:
        return False, None, None, err
    if not text_content or not text_content.strip():
        return False, None, None, "❌ File is empty or no extractable text found."
    base = os.path.basename(path)
    stem, _ext = os.path.splitext(base)
    filename = f"{stem}.txt"
    return upload_text_to_backend(text_content, filename)


def query_backend(message: str, doc_ids=None) -> Generator[str, None, None]:
    """
    POST /query. Streams status, then yields the final answer.

    `doc_ids` may be:
      - None or [] → search across all indexed documents
      - a single filename string → filter to that one doc (legacy)
      - a list of filename strings → filter to any of those docs (OR)
    """
    if not message or not message.strip():
        return
    if len(message) > MAX_QUESTION_LENGTH:
        message = message[:MAX_QUESTION_LENGTH]
    yield "🤔 Thinking..."

    # Normalise doc_ids to a clean list of strings.
    if doc_ids is None:
        id_list = []
    elif isinstance(doc_ids, str):
        id_list = [doc_ids] if doc_ids.strip() else []
    else:
        id_list = [str(d).strip() for d in doc_ids if d and str(d).strip()]

    try:
        payload: dict = {"question": message}
        if id_list:
            payload["doc_ids"] = id_list
        response = requests.post(
            get_backend_url("query"),
            json=payload,
            timeout=QUERY_TIMEOUT,
        )
        response.raise_for_status()
        data = response.json()
        answer = clean_answer(data.get("answer", "⚠️ Received an empty response."))

        # Show which document(s) the answer came from.
        sources = data.get("sources") or []
        if sources:
            src_chips = " ".join(f"`{s}`" for s in sources)
            answer += f"\n\n<sub style='color:#666'>📎 Sources: {src_chips}</sub>"

        elapsed = data.get("elapsed")
        if elapsed is not None:
            answer += f"\n<sub style='color:#888'>⏱️ {elapsed:.1f}s</sub>"
        yield answer
    except requests.exceptions.Timeout:
        yield "❌ Query timed out. Please try again."
    except requests.exceptions.ConnectionError:
        yield "❌ Cannot connect to backend. Please try again later."
    except Exception as e:
        yield f"❌ Error: {str(e)}"


def load_index_state() -> Tuple[bool, list, Optional[str], str]:
    """Page-load helper. Returns (is_indexed, docs, active_doc, display_text)."""
    is_indexed, docs = check_health()
    active = docs[0] if docs else None
    return is_indexed, docs, active, update_index_display(docs)
